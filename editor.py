from PyQt5.QtWidgets import QTextEdit, QToolBar, QAction, QFileDialog
from PyQt5.QtGui import QTextCursor, QTextCharFormat, QFont, QImage
from PyQt5.QtCore import Qt, pyqtSignal, QObject, QTimer
from docx import Document
from docx.shared import Inches
from difflib import SequenceMatcher

class DocumentEditor(QObject):
    change_tracked = pyqtSignal(str)  # Signal to notify when a change is tracked

    def __init__(self, change_tracker):
        super().__init__()
        self.editor_widget = QTextEdit()
        self.editor_widget.setReadOnly(True)  # Disable editing initially
        self.document = None
        self.change_tracker = change_tracker
        self.original_content = ""  # To compare changes
        self.previous_content = ""  # Store the last known content
        self.first_change_logged = False  # Flag to ensure the first change is captured
        self.initial_comparison_done = False  # Flag to handle initial content setup

        # Buffers for aggregating character insertions and deletions
        self.insertion_buffer = ""
        self.deletion_buffer = ""
        self.buffer_timer = QTimer()
        self.buffer_timer.setSingleShot(True)
        self.buffer_timer.timeout.connect(self.flush_buffers)

        # Setup the toolbar
        self.toolbar = QToolBar("Editor Toolbar")
        self.init_toolbar()

        # Connect text change signal
        self.editor_widget.textChanged.connect(self.track_text_change)

    def get_editor_widget(self):
        return self.editor_widget

    def get_toolbar(self):
        return self.toolbar

    def init_toolbar(self):
        bold_action = QAction("Bold", self.toolbar)
        bold_action.triggered.connect(self.make_bold)
        self.toolbar.addAction(bold_action)

        italic_action = QAction("Italic", self.toolbar)
        italic_action.triggered.connect(self.make_italic)
        self.toolbar.addAction(italic_action)

        undo_action = QAction("Undo", self.toolbar)
        undo_action.triggered.connect(self.editor_widget.undo)
        self.toolbar.addAction(undo_action)

        redo_action = QAction("Redo", self.toolbar)
        redo_action.triggered.connect(self.editor_widget.redo)
        self.toolbar.addAction(redo_action)

        insert_image_action = QAction("Insert Image", self.toolbar)
        insert_image_action.triggered.connect(self.insert_image)
        self.toolbar.addAction(insert_image_action)

    def make_bold(self):
        cursor = self.editor_widget.textCursor()
        fmt = QTextCharFormat()
        fmt.setFontWeight(QFont.Bold if not cursor.charFormat().fontWeight() == QFont.Bold else QFont.Normal)
        cursor.mergeCharFormat(fmt)

    def make_italic(self):
        cursor = self.editor_widget.textCursor()
        fmt = QTextCharFormat()
        fmt.setFontItalic(not cursor.charFormat().fontItalic())
        cursor.mergeCharFormat(fmt)

    def load_document(self, path):
        self.document = Document(path)
        self.original_content = self._get_full_text_from_doc()
        self.previous_content = self.original_content
        self.initial_comparison_done = True
        self.display_document_as_html()
        self.editor_widget.setReadOnly(False)

    def _get_full_text_from_doc(self):
        text = ""
        for para in self.document.paragraphs:
            for run in para.runs:
                if run.element.xpath('.//w:drawing'):
                    text += "<img>"  # Placeholder for image
                else:
                    text += run.text
            text += "\n"
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        return text.strip()

    def display_document_as_html(self):
        html_content = "<html><body>"
        for para in self.document.paragraphs:
            para_text = ""
            for run in para.runs:
                if run.element.xpath('.//w:drawing'):
                    para_text += '<img src="data:image/png;base64,...">'  # Add the actual image data
                else:
                    para_text += run.text.strip().replace('<', '&lt;').replace('>', '&gt;')
            if para_text:
                font = para.style.font
                style = "font-family: {0};".format(font.name if font.name else "default")
                if font.size:
                    style += f" font-size: {font.size.pt}px;"
                if font.bold:
                    style += " font-weight: bold;"
                if font.italic:
                    style += " font-style: italic;"
                if font.underline:
                    style += " text-decoration: underline;"
                html_content += f'<p style="{style}">{para_text}</p>'
        for table in self.document.tables:
            html_content += "<table border='1'>"
            for row in table.rows:
                html_content += "<tr>"
                for cell in row.cells:
                    cell_text = cell.text.replace('<', '&lt;').replace('>', '&gt;')
                    html_content += f"<td>{cell_text}</td>"
                html_content += "</tr>"
            html_content += "</table>"
        html_content += "</body></html>"
        self.editor_widget.setHtml(html_content)

    def insert_image(self):
        file_dialog = QFileDialog()
        image_path, _ = file_dialog.getOpenFileName(self.editor_widget, "Insert Image", "", "Images (*.png *.jpg *.bmp *.gif)")
        if image_path:
            cursor = self.editor_widget.textCursor()
            image = QImage(image_path)
            cursor.insertImage(image)

    def track_text_change(self):
        if self.editor_widget.isReadOnly() or not self.initial_comparison_done:
            return  # Ignore changes when the editor is read-only or if initial comparison isn't ready

        current_content = self.editor_widget.toPlainText()
        matcher = SequenceMatcher(None, self.previous_content, current_content)
        insertions = []
        deletions = []

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'insert':
                insertions.append(current_content[j1:j2])
            elif tag == 'delete':
                deletions.append(self.previous_content[i1:i2])

        # Explicitly handle the first change
        if not self.first_change_logged and (insertions or deletions):
            self.first_change_logged = True

        # Filter out any placeholder text or variables (simple check for now)
        filtered_insertions = [text for text in insertions if not self._is_placeholder(text)]
        filtered_deletions = [text for text in deletions if not self._is_placeholder(text)]

        # Aggregate and buffer insertions and deletions
        if filtered_insertions:
            self.insertion_buffer += "".join(filtered_insertions)
            self.buffer_timer.start(300)  # Wait for 300ms before flushing the buffer
        if filtered_deletions:
            self.deletion_buffer += "".join(filtered_deletions)
            self.buffer_timer.start(300)  # Wait for 300ms before flushing the buffer

        self.previous_content = current_content  # Update content

    def _is_placeholder(self, text):
        # Basic check for placeholder variables or template text (can be expanded)
        return text.startswith("<") and text.endswith(">")

    def flush_buffers(self):
        if self.insertion_buffer:
            self.change_tracker.track_change("insert", self.insertion_buffer, "", "")
            self.change_tracked.emit(f"insert: {self.insertion_buffer}")
            self.insertion_buffer = ""
        if self.deletion_buffer:
            self.change_tracker.track_change("delete", self.deletion_buffer, "", "")
            self.change_tracked.emit(f"delete: {self.deletion_buffer}")
            self.deletion_buffer = ""

    def save_document(self, path):
        edited_content = self.editor_widget.toPlainText()
        self.document = Document()
        for line in edited_content.split("\n"):
            self.document.add_paragraph(line)
        self.document.save(path)
